import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Data for the reports
scenarios = [
    {"ID": "TS-01", "Nama": "End-to-End Onboarding", "Prioritas": "HIGH", "Deskripsi": "Penambahan karyawan baru hingga muncul di Terminal Absensi."},
    {"ID": "TS-02", "Nama": "Siklus Kehadiran Harian", "Prioritas": "HIGH", "Deskripsi": "Proses Masuk dan Pulang karyawan via Terminal."},
    {"ID": "TS-03", "Nama": "Koreksi Absensi HR", "Prioritas": "MEDIUM", "Deskripsi": "Admin melakukan koreksi waktu/status absensi yang salah."},
    {"ID": "TS-04", "Nama": "Validasi Payroll Otomatis", "Prioritas": "HIGH", "Deskripsi": "Verifikasi perhitungan THP (Gaji Pokok + Lembur - Potongan)."},
    {"ID": "TS-05", "Nama": "Real-time Dashboard Sync", "Prioritas": "MEDIUM", "Deskripsi": "Perubahan data absensi langsung mengupdate widget statistik."},
    {"ID": "TS-06", "Nama": "Audit Trail & Data Cleanse", "Prioritas": "LOW", "Deskripsi": "Ekspor data ke Excel dan pembersihan log sistem."},
    {"ID": "TS-07", "Nama": "Keamanan Akses PIN", "Prioritas": "HIGH", "Deskripsi": "Validasi akses terminal hanya untuk karyawan pemilik PIN benar."},
    {"ID": "TS-08", "Nama": "Validasi Input Kritis", "Prioritas": "HIGH", "Deskripsi": "Penanganan input NIK duplikat atau format nominal gaji salah."}
]

test_cases = [
    {"ID": "TC-01.1", "Scenario": "TS-01", "Langkah": "Klik '+ Tambah Karyawan' di Admin", "Input": "-", "Hasil": "Modal input muncul.", "Trace": "Admin UI"},
    {"ID": "TC-01.2", "Scenario": "TS-01", "Langkah": "Isi form lengkap & Klik 'Simpan'", "Input": "Budi Santoso, NIK: 3275..., PIN: 1234", "Hasil": "Notifikasi sukses, data tersimpan.", "Trace": "Database"},
    {"ID": "TC-01.3", "Scenario": "TS-01", "Langkah": "Refresh Terminal Absensi", "Input": "-", "Hasil": "Nama Budi muncul di Terminal.", "Trace": "Terminal UI"},
    {"ID": "TC-04.1", "Scenario": "TS-04", "Langkah": "Isi Gaji Pokok karyawan", "Input": "Rp 5.000.000", "Hasil": "Hitung Gaji Proporsional.", "Trace": "Logic (admin.js)"},
    {"ID": "TC-04.2", "Scenario": "TS-04", "Langkah": "Input Jam Lembur", "Input": "10 Jam", "Hasil": "Uang Lembur = 10 * Tarif.", "Trace": "Logic (admin.js)"},
    {"ID": "TC-04.3", "Scenario": "TS-04", "Langkah": "Verifikasi Potongan", "Input": "-", "Hasil": "THP berkurang otomatis by BPJS & PPh21.", "Trace": "Audit Payroll"}
]

rtm = [
    {"ReqID": "REQ-01", "Fitur": "Master Data Karyawan", "Scenario": "TS-01, TS-06", "TestCases": "TC-01.1, TC-01.2", "Status": "Ready"},
    {"ReqID": "REQ-02", "Fitur": "Terminal Absensi", "Scenario": "TS-02, TS-07", "TestCases": "TC-02.1, TC-07.1", "Status": "Ready"},
    {"ReqID": "REQ-03", "Fitur": "Koreksi Data HR", "Scenario": "TS-03", "TestCases": "TC-03.1, TC-03.2", "Status": "Ready"},
    {"ReqID": "REQ-04", "Fitur": "Kalkulasi Payroll", "Scenario": "TS-04", "TestCases": "TC-04.1, TC-04.2, TC-04.3", "Status": "Ready"},
    {"ReqID": "REQ-05", "Fitur": "Dashboard Stats", "Scenario": "TS-05", "TestCases": "TC-05.1", "Status": "Ready"},
    {"ReqID": "REQ-06", "Fitur": "Excel Reporting", "Scenario": "TS-06", "TestCases": "TC-06.1", "Status": "Ready"}
]

def generate_excel():
    with pd.ExcelWriter('KOBOI_Testing_Report.xlsx', engine='openpyxl') as writer:
        pd.DataFrame(scenarios).to_excel(writer, sheet_name='Scenarios', index=False)
        pd.DataFrame(test_cases).to_excel(writer, sheet_name='Test Cases', index=False)
        pd.DataFrame(rtm).to_excel(writer, sheet_name='Traceability Matrix', index=False)
    print("Excel report generated: KOBOI_Testing_Report.xlsx")

def generate_word():
    doc = Document()
    
    # Title
    title = doc.add_heading('Laporan Pengujian Sistem Informasi KOBOI (HRIS)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Section 1: Strategic Test Plan
    doc.add_heading('1. Strategic Test Plan', level=1)
    doc.add_paragraph('Objektif: Memastikan integritas data payroll dan keakuratan log absensi real-time.')
    
    # Section 2: Scenarios
    doc.add_heading('2. Business Scenarios', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ID'
    hdr_cells[1].text = 'Nama Skenario'
    hdr_cells[2].text = 'Prioritas'
    
    for s in scenarios:
        row_cells = table.add_row().cells
        row_cells[0].text = s['ID']
        row_cells[1].text = s['Nama']
        row_cells[2].text = s['Prioritas']
        
    # Section 3: Test Cases
    doc.add_heading('3. Detailed Test Cases', level=1)
    for tc in test_cases:
        p = doc.add_paragraph()
        p.add_run(f"{tc['ID']}: {tc['Langkah']}").bold = True
        doc.add_paragraph(f"Input: {tc['Input']}")
        doc.add_paragraph(f"Hasil yang Diharapkan: {tc['Hasil']}")
        doc.add_paragraph(f"Traceability: {tc['Trace']}")
        doc.add_paragraph("-" * 20)

    doc.save('KOBOI_Testing_Report.docx')
    print("Word report generated: KOBOI_Testing_Report.docx")

if __name__ == "__main__":
    generate_excel()
    generate_word()
